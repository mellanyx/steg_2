import docx
import MTK2
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document('text_clear.docx')

def run_set_spacing(run, value: int):
    """Set the font spacing for `run` to `value` in twips.

    A twip is a "twentieth of an imperial point", so 1/1440 in.
    """

    def get_or_add_spacing(rPr):
        # --- check if `w:spacing` child already exists ---
        spacings = rPr.xpath("./w:spacing")
        # --- return that if so ---
        if spacings:
            return spacings[0]
        # --- otherwise create one ---
        spacing = OxmlElement("w:spacing")
        rPr.insert_element_before(
            spacing,
            *(
                "w:w",
                "w:kern",
                "w:position",
                "w:sz",
                "w:szCs",
                "w:highlight",
                "w:u",
                "w:effect",
                "w:bdr",
                "w:shd",
                "w:fitText",
                "w:vertAlign",
                "w:rtl",
                "w:cs",
                "w:em",
                "w:lang",
                "w:eastAsianLayout",
                "w:specVanish",
                "w:oMath",
            ),
        )
        return spacing

    rPr = run._r.get_or_add_rPr()
    spacing = get_or_add_spacing(rPr)
    spacing.set(qn('w:val'), str(value))
    # spacing.set("val", str(value))

if __name__ == '__main__':
    text = "Жизнь прожить - не поле перейти."

    TextMTK2 = MTK2.MTK2_code(text)
    print(TextMTK2)

    LenParagraphs = []
    OpenText = ""
    for paragraph in doc.paragraphs:
        stroka = ""
        for run in paragraph.runs:
            for char in run.text:
                stroka += char
        OpenText += stroka
        LenParagraphs.append(len(stroka))

    print(OpenText)
    print(LenParagraphs)

    doc.paragraphs.clear()
    id_char = 0
    for id_paragraph in range(len(doc.paragraphs)):
        doc.paragraphs[id_paragraph].clear()
        for id_rans in range(LenParagraphs[id_paragraph]):
            run = doc.paragraphs[id_paragraph].add_run(OpenText[id_char])
            if id_char < len(TextMTK2):
                if TextMTK2[id_char] == '1':
                    run_set_spacing(run, -10)
                    # run.font.size = Pt(12)
                    # run.font.color.rgb = RGBColor(0, 0, 1)
                else:
                    run_set_spacing(run, 0)
                    # run.font.size = Pt(11)
                    # run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run_set_spacing(run, 0)
                # run.font.size = Pt(11)
                # run.font.color.rgb = RGBColor(0, 0, 0)
            id_char += 1

    doc.save('text_clear_save.docx')
    print("Файл text_clear_save.docx изменён и сохранён!")